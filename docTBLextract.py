#!/usr/bin/env python3
"""
       /$$                  /$$$$$$$$ /$$$$$$$  /$$                             /$$                                    /$$    
      | $$                 |__  $$__/| $$__  $$| $$                            | $$                                   | $$    
  /$$$$$$$  /$$$$$$   /$$$$$$$| $$   | $$  \ $$| $$        /$$$$$$  /$$   /$$ /$$$$$$    /$$$$$$  /$$$$$$   /$$$$$$$ /$$$$$$  
 /$$__  $$ /$$__  $$ /$$_____/| $$   | $$$$$$$ | $$       /$$__  $$|  $$ /$$/|_  $$_/   /$$__  $$|____  $$ /$$_____/|_  $$_/  
| $$  | $$| $$  \ $$| $$      | $$   | $$__  $$| $$      | $$$$$$$$ \  $$$$/   | $$    | $$  \__/ /$$$$$$$| $$        | $$    
| $$  | $$| $$  | $$| $$      | $$   | $$  \ $$| $$      | $$_____/  >$$  $$   | $$ /$$| $$      /$$__  $$| $$        | $$ /$$
|  $$$$$$$|  $$$$$$/|  $$$$$$$| $$   | $$$$$$$/| $$$$$$$$|  $$$$$$$ /$$/\  $$  |  $$$$/| $$     |  $$$$$$$|  $$$$$$$  |  $$$$/
 \_______/ \______/  \_______/|__/   |_______/ |________/ \_______/|__/  \__/   \___/  |__/      \_______/ \_______/   \___/  

Convert tables from Microsoft Word docs to Excel
-
Author:
sorzkode
https://github.com/sorzkode

MIT License
Copyright (c) 2025
"""

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
import random
import string
from docx import Document
from openpyxl import Workbook

class WordToExcelConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("docTBLextract")
        self.root.geometry("600x400")
        self.root.iconbitmap("assets/icon.ico")
        
        # Variables
        self.word_file_path = None
        self.table_count = 0
        
        # Style
        style = ttk.Style()
        style.theme_use('clam')
        
        # Main frame
        main_frame = ttk.Frame(root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights
        root.columnconfigure(0, weight=1)
        root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Title
        title_label = ttk.Label(main_frame, text="Convert Tables From Word To Excel", 
                               font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 20))
        
        # Word file selection
        ttk.Label(main_frame, text="Word Document:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.file_label = ttk.Label(main_frame, text="No file selected", 
                                   relief=tk.SUNKEN, anchor=tk.W)
        self.file_label.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=10, pady=5)
        
        self.browse_button = ttk.Button(main_frame, text="Browse", 
                                       command=self.browse_word_file)
        self.browse_button.grid(row=1, column=2, pady=5)
        
        # Table info
        self.info_frame = ttk.LabelFrame(main_frame, text="Document Information", 
                                        padding="10")
        self.info_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), 
                            pady=20)
        
        self.table_info_label = ttk.Label(self.info_frame, 
                                         text="Please select a Word document")
        self.table_info_label.grid(row=0, column=0)
        
        # Convert button
        self.convert_button = ttk.Button(main_frame, text="Convert to Excel", 
                                        command=self.convert_to_excel, 
                                        state=tk.DISABLED)
        self.convert_button.grid(row=3, column=0, columnspan=3, pady=20)
        
        # Progress bar
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), 
                          pady=(0, 10))
        
        # Status label
        self.status_label = ttk.Label(main_frame, text="Ready", 
                                     foreground="green")
        self.status_label.grid(row=5, column=0, columnspan=3)
        
    def browse_word_file(self):
        """Open file dialog to select a Word document"""
        filename = filedialog.askopenfilename(
            title="Select Word Document",
            filetypes=[("Word Documents", "*.docx"), ("All Files", "*.*")]
        )
        
        if filename:
            self.word_file_path = filename
            self.file_label.config(text=os.path.basename(filename))
            self.analyze_document()
            
    def analyze_document(self):
        """Analyze the Word document for tables"""
        try:
            self.status_label.config(text="Analyzing document...", 
                                   foreground="blue")
            self.root.update()
            
            doc = Document(self.word_file_path)
            self.table_count = len(doc.tables)
            
            if self.table_count == 0:
                self.table_info_label.config(
                    text="⚠️ No tables found in the document",
                    foreground="red"
                )
                self.convert_button.config(state=tk.DISABLED)
                messagebox.showwarning(
                    "No Tables Found",
                    "The selected Word document does not contain any tables."
                )
            else:
                self.table_info_label.config(
                    text=f"✓ Found {self.table_count} table(s) in the document",
                    foreground="green"
                )
                self.convert_button.config(state=tk.NORMAL)
                
            self.status_label.config(text="Ready", foreground="green")
            
        except Exception as e:
            self.status_label.config(text="Error analyzing document", 
                                   foreground="red")
            messagebox.showerror("Error", f"Failed to analyze document:\n{str(e)}")
            
    def generate_random_name(self, length=8):
        """Generate a random worksheet name"""
        return ''.join(random.choices(string.ascii_letters + string.digits, 
                                    k=length))
        
    def convert_to_excel(self):
        """Convert Word tables to Excel worksheets"""
        # Ask for save location
        save_path = filedialog.asksaveasfilename(
            title="Save Excel File As",
            defaultextension=".xlsx",
            filetypes=[("Excel Files", "*.xlsx"), ("All Files", "*.*")]
        )
        
        if not save_path:
            return
            
        try:
            # Start progress bar
            self.progress.start()
            self.status_label.config(text="Converting tables...", 
                                   foreground="blue")
            self.convert_button.config(state=tk.DISABLED)
            self.root.update()
            
            # Read Word document
            doc = Document(self.word_file_path)
            
            # Create Excel workbook
            wb = Workbook()
            wb.remove(wb.active)  # Remove default sheet
            
            # Process each table
            for idx, table in enumerate(doc.tables):
                # Create worksheet with random name
                ws_name = f"Table_{self.generate_random_name()}"
                ws = wb.create_sheet(title=ws_name)
                
                # Copy table data
                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        # Excel uses 1-based indexing
                        ws.cell(row=row_idx + 1, column=col_idx + 1, 
                               value=cell.text)
                        
                # Auto-adjust column widths
                for column in ws.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                            
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column_letter].width = adjusted_width
                    
            # Save Excel file
            wb.save(save_path)
            
            # Stop progress bar
            self.progress.stop()
            self.status_label.config(text="Conversion completed successfully!", 
                                   foreground="green")
            self.convert_button.config(state=tk.NORMAL)
            
            messagebox.showinfo(
                "Success",
                f"Successfully converted {self.table_count} table(s) to Excel!\n"
                f"Saved to: {os.path.basename(save_path)}"
            )
            
        except Exception as e:
            self.progress.stop()
            self.status_label.config(text="Conversion failed", 
                                   foreground="red")
            self.convert_button.config(state=tk.NORMAL)
            messagebox.showerror("Error", f"Failed to convert:\n{str(e)}")

def main():
    root = tk.Tk()
    app = WordToExcelConverter(root)
    root.mainloop()

if __name__ == "__main__":
    main()